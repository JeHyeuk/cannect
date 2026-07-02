from cannect.config import env
from cannect.core.can.db.reader import CANDBReader
from cannect.core.can.db.specification.styles import CustomStyle
from cannect.core.can.db.specification.message import Message
from cannect.core.subversion import Subversion
from cannect.utils import tools

from datetime import datetime
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches
from pandas import DataFrame
from pathlib import Path
from tqdm.auto import tqdm
from typing import Union
import os, time, site


class Specification:

    def __init__(self, db:Union[CANDBReader, DataFrame]):
        if isinstance(db, DataFrame):
            db = CANDBReader(db)
        self._db = db

        if not (env.SVN_CANDB / r"사양서/resource").exists():
            Subversion.update(env.SVN_CANDB / "사양서")

        template = Path(site.getsitepackages()[1]) / 'docx/templates/default.docx'
        if template.exists():
            os.remove(template)
        tools.copy_to(env.SVN_CANDB / r"사양서/resource/default", template.parent)
        time.sleep(0.5)
        os.rename(template.parent / 'default', template.parent / 'default.docx')
        time.sleep(0.5)

        self._doc = doc = Document()
        self._style = CustomStyle(doc)
        return

    def _set_title(self):
        self._doc.add_paragraph(
            "\n자체제어기 EMS/ASW\n통신 사양서(CAN-FD)\n\n",
            style=self._style.title
        )
        return

    def _set_ci(self):
        png = env.SVN_CANDB / r"사양서/resource/ci_cover.png"
        paragraph = self._doc.add_paragraph()
        runner = paragraph.add_run()
        runner.add_picture(str(png), width=Inches(6))
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        self._doc.add_paragraph("\n", style=self._style.title)
        return

    def _set_overview(self):
        items = {
            'DATABASE': f"자체제어기/EMS CAN-FD {self._db.revision}",
            'COMPANY': env.COMPANY,
            'DIVISION': env.DIVISION,
            'RELEASE': datetime.now().strftime("%Y-%m-%d"),
        }

        table = self._doc.add_table(rows=len(items), cols=2)
        table.style = 'Table Grid'
        for n, (key, value) in enumerate(items.items()):
            left = table.rows[n].cells[0]
            left.width = Inches(1)
            name = left.paragraphs[0]
            name.style = self._style.overview_left
            name.text = key

            right = table.rows[n].cells[1]
            right.width = self._doc.sections[0].page_width
            text = right.paragraphs[0]
            text.style = self._style.overview_right
            text.text = value
        return

    def _set_margin(self):
        section = self._doc.sections[0]
        section.left_margin = \
        section.right_margin = Inches(0.5)
        section.bottom_margin = \
        section.top_margin = Inches(1)
        return

    def _set_header(self, version:str):
        header = self._doc.sections[0].header
        for paragraph in header.paragraphs:
            p = getattr(paragraph, "_element")
            p.getparent().remove(p)

        table = header.add_table(rows=1, cols=3, width=self._doc.sections[0].page_width)

        left = table.rows[0].cells[0].paragraphs[0]
        left.text = "자체제어기 EMS/ASW CAN-FD" + "\n" + f"DOC {version}"
        left.style = self._style.header_left

        right = table.rows[0].cells[2].paragraphs[0]
        right.text = env.DIVISION + "\n" + env.COMPANY
        right.style = self._style.header_right
        return

    def _set_footer(self):
        footer = self._doc.sections[0].footer
        for paragraph in footer.paragraphs:
            p = getattr(paragraph, "_element")
            p.getparent().remove(p)
        table = footer.add_table(rows=1, cols=1, width=self._doc.sections[0].page_width)
        cell = table.rows[0].cells[0].paragraphs[0]
        cell.text = env.COPYRIGHT
        cell.style = self._style.footer
        return

    def generate(self, filename:str):
        if not filename.endswith('.docx'):
            filename += '.docx'
        objs = [(msg, obj) for msg, obj in self._db.messages.items()]
        objs = sorted(objs, key=lambda x: x[0])

        self._set_margin()
        self._set_header(self._db.revision)
        self._set_footer()
        self._set_title()
        self._set_ci()
        self._set_overview()

        self._doc.add_section()
        message = Message(self._doc)
        message.addHeading("EMS TRANSMIT")
        transmit = tqdm([obj for _, obj in objs if obj["ECU"] == "EMS"])
        for obj in transmit:
            transmit.set_description(desc=f"{obj.name} 사양 생성")
            message.addMessageHeading(obj)
            message.addMessageSpec(obj)
            message.addMessageLayout(obj)
            message.addSignalList(obj)
            message.addSignalProperty(obj)

        message.addHeading("EMS RECEIVE")
        receive = tqdm([obj for _, obj in objs if obj["ECU"] != "EMS"])
        for obj in receive:
            receive.set_description(desc=f"{obj.name} 사양 생성")
            message.addMessageHeading(obj)
            message.addMessageSpec(obj)
            message.addMessageLayout(obj)
            message.addSignalList(obj)
            message.addSignalProperty(obj)

        self._doc.save(env.DOWNLOADS / filename)
        return


if __name__ == "__main__":
    db = CANDBReader().mode('HEV')

    spec = Specification(db)
    spec.generate("TEST CAN SPEC")